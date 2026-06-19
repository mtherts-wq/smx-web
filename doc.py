import customtkinter as ctk
from tkinter import filedialog, messagebox
from docx import Document
from datetime import datetime
from tkcalendar import DateEntry
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from PIL import Image
import os, sys
import win32com.client

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

fotos = []
caminho_docx = None

# ========================
# MODELO OCULTO (CRÍTICO)
# ========================

def caminho_modelo():
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, "MODELO_RAT.docx")
    return "MODELO_RAT.docx"

# ========================
# VALIDAÇÃO
# ========================

def validar_hora(valor):
    try:
        datetime.strptime(valor, "%H:%M")
        return True
    except:
        return False

def marcar_erro(campo):
    campo.configure(border_color="red", border_width=2)

def limpar_erro(campo):
    campo.configure(border_width=0)

def validar_campos():
    erro = False

    if not validar_hora(e_inicio.get()):
        marcar_erro(e_inicio)
        erro = True
    else:
        limpar_erro(e_inicio)

    if not validar_hora(e_fim.get()):
        marcar_erro(e_fim)
        erro = True
    else:
        limpar_erro(e_fim)

    if erro:
        messagebox.showerror("Erro", "Horário inválido HH:MM")
        return False

    ini = datetime.strptime(e_inicio.get(), "%H:%M")
    fim = datetime.strptime(e_fim.get(), "%H:%M")

    if fim <= ini:
        messagebox.showerror("Erro", "Hora fim deve ser maior que início")
        return False

    return True

# ========================
# DADOS
# ========================

def calcular_tempo(inicio, fim):
    t1 = datetime.strptime(inicio, "%H:%M")
    t2 = datetime.strptime(fim, "%H:%M")
    diff = t2 - t1
    h, r = divmod(diff.total_seconds(), 3600)
    m = r // 60
    return f"{int(h):02}:{int(m):02}"

def gerar_dados():
    return {
        "{{PROTOCOLO}}": e_protocolo.get(),
        "{{TITULO}}": e_titulo.get(),
        "{{ATENDENTE}}": e_atendente.get(),
        "{{LOJA}}": combo_loja.get(),
        "{{LOCAL}}": e_local.get(),
        "{{TECNICO}}": e_tecnico.get(),
        "{{EMPRESA}}": "SMX",
        "{{DATA}}": data_picker.get(),
        "{{HORARIO}}": f"{e_inicio.get()} as {e_fim.get()}",
        "{{TEMPO}}": calcular_tempo(e_inicio.get(), e_fim.get()),
        "{{GERENTE}}": e_gerente.get(),
        "{{DESCRICAO}}": t_desc.get("0.0", "end").strip()
    }

# ========================
# ENGINE UNIVERSAL ✅
# ========================

def substituir_tudo(doc, dados):

    def processar(p):
        texto = "".join(run.text for run in p.runs)

        if "{{" in texto:
            for chave, valor in dados.items():
                texto = texto.replace(chave, valor)

            p.clear()
            run = p.add_run(texto)
            run.font.name = "Calibri"
            run.font.size = Pt(11)

            if "–" in texto:
                run.bold = True

    for p in doc.paragraphs:
        processar(p)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    valor_original = p.text
                    processar(p)
                    if any(tag in valor_original for tag in dados):
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# ========================
# FOTOS
# ========================

def inserir_fotos(doc):
    for p in doc.paragraphs:
        if "Validado com" in p.text:

            if fotos:
                total = len(fotos)
                linhas = (total // 2) + (total % 2)

                tabela = doc.add_table(rows=linhas, cols=2)
                tabela.alignment = WD_TABLE_ALIGNMENT.CENTER

                idx = 0

                for r in range(linhas):
                    for c in range(2):
                        if idx < total:

                            img = Image.open(fotos[idx])
                            w, h = img.size
                            ratio = w / h

                            if ratio > 1:
                                largura, altura = 5, 5/ratio
                            else:
                                altura, largura = 8, 8*ratio

                            cell = tabela.cell(r, c)
                            run = cell.paragraphs[0].add_run()

                            run.add_picture(
                                fotos[idx],
                                width=Cm(largura),
                                height=Cm(altura)
                            )

                            idx += 1

                p._element.addprevious(tabela._element)

                for _ in range(5):
                    p.insert_paragraph_before()

            break

# ========================
# GERAR DOCUMENTO
# ========================

def gerar_documento():
    global caminho_docx

    try:
        if not validar_campos():
            return

        doc = Document(caminho_modelo())

        dados = gerar_dados()

        substituir_tudo(doc, dados)
        inserir_fotos(doc)

        data_fmt = datetime.strptime(
            data_picker.get(), "%d/%m/%Y"
        ).strftime("%d%m")

        loja = combo_loja.get().upper()

        base = e_local.get() if loja == "ZARA" else combo_loja.get()

        nome = f"{base}_{data_fmt}.docx"
        caminho_docx = os.path.abspath(nome)

        doc.save(caminho_docx)

        messagebox.showinfo("Sucesso", f"Gerado:\n{nome}")
        os.startfile(caminho_docx)

    except Exception as e:
        messagebox.showerror("Erro", str(e))

# ========================
# PDF
# ========================

def abrir_pdf():
    if not caminho_docx:
        messagebox.showwarning("Aviso", "Gere primeiro o documento")
        return

    word = win32com.client.Dispatch("Word.Application")
    doc = word.Documents.Open(caminho_docx)

    pdf = caminho_docx.replace(".docx", ".pdf")

    doc.SaveAs(pdf, FileFormat=17)
    doc.Close()
    word.Quit()

    os.startfile(pdf)

# ========================
# RESET
# ========================

def resetar():
    for c in [e_protocolo,e_titulo,e_atendente,
              e_local,e_tecnico,e_inicio,
              e_fim,e_gerente]:
        c.delete(0,"end")

    t_desc.delete("0.0","end")
    fotos.clear()

# ========================
# UI
# ========================

app = ctk.CTk()
app.geometry("620x850")
app.title("SMX Gerador")

def campo(label):
    ctk.CTkLabel(app,text=label).pack()
    e = ctk.CTkEntry(app,width=350)
    e.pack()
    return e

e_protocolo = campo("Protocolo")
e_titulo = campo("Título")
e_atendente = campo("Atendente SMX")

ctk.CTkLabel(app,text="Loja").pack()
combo_loja = ctk.CTkComboBox(app,values=[
    "ZARA","Clinica CLIMAR","Clinica Marcelo Araújo",
    "DALGAS Escritório","DALGAS G7","VIVARA"
])
combo_loja.pack()

e_local = campo("Local")
e_tecnico = campo("Técnico")

ctk.CTkLabel(app,text="Data").pack()
data_picker = DateEntry(app,date_pattern="dd/mm/yyyy")
data_picker.pack()

e_inicio = campo("Hora início")
e_fim = campo("Hora fim")
e_gerente = campo("Gerente")

ctk.CTkLabel(app,text="Descrição").pack()
t_desc = ctk.CTkTextbox(app,width=420,height=120)
t_desc.pack()

frame = ctk.CTkFrame(app)
frame.pack(pady=20)

ctk.CTkButton(frame,text="Fotos",command=selecionar_fotos)\
.grid(row=0,column=0,padx=5)

ctk.CTkButton(frame,text="Gerar",command=gerar_documento)\
.grid(row=0,column=1,padx=5)

ctk.CTkButton(frame,text="Abrir PDF",command=abrir_pdf)\
.grid(row=0,column=2,padx=5)

ctk.CTkButton(app,text="Reset",command=resetar,width=200).pack(pady=10)

app.mainloop()