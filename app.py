from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.shared import Cm
from datetime import datetime
import os
import pandas as pd
import subprocess
import uuid
import threading

app = Flask(__name__)

# ========================
# CONTROLE DE PROGRESSO
# ========================
# Armazenará um dicionário por job_id: {"progresso": int, "path": str, "nome": str}
progresso = {}
ARQUIVO_HISTORICO = "historico.xlsx"

# ========================
# TEMPO
# ========================
def calcular_tempo(inicio, fim):
    try:
        t1 = datetime.strptime(inicio, "%H:%M")
        t2 = datetime.strptime(fim, "%H:%M")
        diff = t2 - t1
        h = diff.seconds // 3600
        m = (diff.seconds % 3600) // 60
        return f"{h:02}:{m:02}"
    except:
        return "00:00"

# ========================
# VALIDAÇÃO
# ========================
def validar_campos(form):
    campos = [
        "protocolo", "titulo", "atendente", "loja",
        "local", "tecnico", "data", "inicio",
        "fim", "gerente", "descricao"
    ]
    for c in campos:
        if not form.get(c):
            return f"Campo obrigatório: {c}"
    return None

# ========================
# NOME (REGRA ZARA)
# ========================
def gerar_nome(dados):
    data = dados.get("{{DATA}}", "")
    loja = (dados.get("{{LOJA}}") or "").strip()
    local = (dados.get("{{LOCAL}}") or "").strip()

    try:
        dt = datetime.strptime(data, "%d/%m/%Y")
        data_fmt = dt.strftime("%d%m")
    except:
        data_fmt = "0000"

    base = local if loja.upper() == "ZARA" else loja
    base = (base or "arquivo").replace(" ", "_")

    return f"{base}_{data_fmt}"

# ========================
# HISTÓRICO
# ========================
def salvar_historico(dados):
    registro = {
        "Data": dados.get("{{DATA}}"),
        "Título": dados.get("{{TITULO}}"),
        "Loja": dados.get("{{LOJA}}"),
        "Atendente": dados.get("{{ATENDENTE}}"),
        "Local": dados.get("{{LOCAL}}"),
        "Tempo": dados.get("{{TEMPO}}"),
        "Gerente": dados.get("{{GERENTE}}")
    }

    df_new = pd.DataFrame([registro])

    if os.path.exists(ARQUIVO_HISTORICO):
        try:
            df_old = pd.read_excel(ARQUIVO_HISTORICO)
            df = pd.concat([df_old, df_new], ignore_index=True)
        except:
            df = df_new
    else:
        df = df_new

    df.to_excel(ARQUIVO_HISTORICO, index=False)

# ========================
# SUBSTITUIÇÃO (TEXTO E TABELAS)
# ========================
def substituir_campos(doc, dados):
    # Substitui nos parágrafos comuns
    for p in doc.paragraphs:
        for k, v in dados.items():
            if k in p.text:
                for run in p.runs:
                    run.text = run.text.replace(k, str(v or ""))

    # Substitui dentro de tabelas
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    for k, v in dados.items():
                        if k in p.text:
                            for run in p.runs:
                                run.text = run.text.replace(k, str(v or ""))

# ========================
# DOCX
# ========================
def gerar_doc(job_id, dados, fotos):
    doc = Document("MODELO_RAT.docx")
    substituir_campos(doc, dados)

    for i, p in enumerate(doc.paragraphs):
        if "Validado com a Gerente" in p.text:
            table = doc.add_table(rows=0, cols=2)

            for idx, f in enumerate(fotos):
                if idx % 2 == 0:
                    row = table.add_row().cells
                if os.path.exists(f):
                    cell = row[idx % 2]
                    run = cell.paragraphs[0].add_run()
                    run.add_picture(f, width=Cm(5), height=Cm(8))

            doc.element.body.insert(i, table._element)
            break

    os.makedirs("temp", exist_ok=True)
    path = f"temp/saida_{job_id}.docx"
    doc.save(path)

    return path

# ========================
# PDF (LibreOffice)
# ========================
def converter_pdf(doc_path):
    try:
        subprocess.run([
            "soffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", "temp",
            doc_path
        ], check=True, timeout=60)
    except Exception as e:
        raise Exception(f"Erro no LibreOffice: {str(e)}")

    return doc_path.replace(".docx", ".pdf")

# ========================
# PROCESSAMENTO ASSÍNCRONO
# ========================
def processar_pdf(job_id, dados, fotos):
    try:
        progresso[job_id]["status"] = 30  # Início do processamento interno

        # 1. Salva no histórico do Excel
        salvar_historico(dados)
        progresso[job_id]["status"] = 50

        # 2. Gera o arquivo .docx customizado
        doc_path = gerar_doc(job_id, dados, fotos)
        progresso[job_id]["status"] = 70

        # 3. Converte para PDF usando LibreOffice
        pdf_path = converter_pdf(doc_path)

        # Remove o arquivo .docx temporário para economizar espaço
        if os.path.exists(doc_path):
            os.remove(doc_path)

        # 4. Finaliza com sucesso guardando o caminho final
        progresso[job_id]["status"] = "concluido"
        progresso[job_id]["path"] = pdf_path

    except Exception as e:
        progresso[job_id]["status"] = f"erro: {str(e)}"

# ========================
# ROTAS
# ========================
@app.route("/")
def home():
    return render_template("index.html")

@app.route("/pdf/start", methods=["POST"])
def iniciar_pdf():
    erro = validar_campos(request.form)
    if erro:
        return erro, 400

    # Captura os dados no Contexto da Requisição ATUAL (evita quebrar na thread)
    data_fmt = datetime.strptime(
        request.form.get("data"), "%Y-%m-%d"
    ).strftime("%d/%m/%Y")

    inicio = request.form.get("inicio")
    fim = request.form.get("fim")

    dados = {
        "{{PROTOCOLO}}": request.form.get("protocolo"),
        "{{TITULO}}": request.form.get("titulo"),
        "{{ATENDENTE}}": request.form.get("atendente"),
        "{{LOJA}}": request.form.get("loja"),
        "{{LOCAL}}": request.form.get("local"),
        "{{TECNICO}}": request.form.get("tecnico"),
        "{{DATA}}": data_fmt,
        "{{HORARIO}}": f"{inicio} as {fim}",
        "{{TEMPO}}": calcular_tempo(inicio, fim),
        "{{GERENTE}}": request.form.get("gerente"),
        "{{DESCRICAO}}": request.form.get("descricao"),
    }

    # Salva os arquivos de foto antes de despachar para a thread
    fotos = []
    os.makedirs("temp", exist_ok=True)
    for f in request.files.getlist("fotos"):
        if f.filename:
            # Garante nome único para evitar colisões
            nome_foto = f"{uuid.uuid4()}_{f.filename}"
            path_foto = os.path.join("temp", nome_foto)
            f.save(path_foto)
            fotos.append(path_foto)

    # Identificadores Únicos
    job_id = str(uuid.uuid4())
    nome_final_arquivo = gerar_nome(dados)

    progresso[job_id] = {
        "status": 10,
        "path": None,
        "nome": nome_final_arquivo
    }

    # Despacha o processamento pesado em segundo plano com dados limpos
    threading.Thread(
        target=processar_pdf,
        args=(job_id, dados, fotos)
    ).start()

    return jsonify({"job_id": job_id})

@app.route("/pdf/status/<job_id>")
def status_pdf(job_id):
    job = progresso.get(job_id)
    if not job:
        return jsonify({"status": "não encontrado"}), 404
    return jsonify({"status": job["status"]})

@app.route("/pdf/download/<job_id>")
def download_pdf(job_id):
    job = progresso.get(job_id)
    if not job or job["status"] != "concluido":
        return "Arquivo ainda não processado ou inválido", 400

    path = job["path"]
    nome_download = f"{job['nome']}.pdf"

    return send_file(path, as_attachment=True, download_name=nome_download)

# ==================================================
# COMO IMPLEMENTAR O RELATÓRIO MENSAL DO EXCEL
# ==================================================
@app.route("/excel")
def excel():
    if not os.path.exists(ARQUIVO_HISTORICO):
        return "Sem registros no histórico ainda", 400

    # Captura o parâmetro de mês enviado pela URL (ex: /excel?mes=06/2026)
    # Se não for enviado, assume o mês e ano corrente baseado no dia de hoje
    mes_filtro = request.args.get("mes")
    if not mes_filtro:
        mes_filtro = datetime.now().strftime("%m/%Y")

    try:
        df = pd.read_excel(ARQUIVO_HISTORICO)
        
        # Garante que a coluna 'Data' seja interpretada como texto para o filtro
        df['Data'] = df['Data'].astype(str)

        # Filtra as linhas onde o final da string corresponde a "MM/AAAA"
        df_filtrado = df[df['Data'].str.endswith(mes_filtro)]

        if df_filtrado.empty:
            return f"Nenhum registro encontrado para o período: {mes_filtro}", 404

        os.makedirs("temp", exist_ok=True)
        path_relatorio = f"temp/relatorio_{mes_filtro.replace('/', '_')}.xlsx"

        # Exporta preservando as exatas colunas que você listou
        df_filtrado.to_excel(path_relatorio, index=False)
        
        return send_file(
            path_relatorio, 
            as_attachment=True, 
            download_name=f"relatorio_{mes_filtro.replace('/', '_')}.xlsx"
        )
    except Exception as e:
        return f"Erro ao gerar relatório: {str(e)}", 500

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)